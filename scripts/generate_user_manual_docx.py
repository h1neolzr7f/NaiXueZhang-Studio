from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "aitag-mirror 使用手册.docx"
FONT = "Microsoft YaHei"


def set_font(run, *, size: float | None = None, bold: bool = False, color=None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=9.5, bold=bold)


def add_kv_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i], "D9EAF7")
        cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)


def paragraph(doc: Document, text: str = ""):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_font(run, size=10.5)
    return p


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = FONT
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 86, 108)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("aitag-mirror 本地图库测试包使用手册")
    set_font(run, size=20, bold=True, color=(31, 78, 121))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("离线 aitag.win / 明日方舟 NovelAI 咒语图库 + 本地 NAI 工作台")
    set_font(run, size=10.5, color=(85, 85, 85))

    paragraph(
        doc,
        "本文档面向收到测试包的用户。发布包包含基础图库数据，可以离线浏览、搜索作品和查看咒语；外部 API token 不随包提供，需用户自行配置。",
    )

    doc.add_heading("1. 一键启动", level=1)
    for step in (
        r"本包已解压到 E:\aitag-mirror-gallery，无需再解压。",
        "双击 START_GALLERY.bat（或桌面快捷方式 aitag-mirror Gallery）。首次启动会自动创建 .venv 并安装 Python 依赖，时间取决于网络。",
        "启动成功后会自动打开 http://127.0.0.1:8797/。如果没有自动打开，可手动复制地址到浏览器。",
        "需要重载代码时运行 start_gallery.bat restart；需要守护进程时运行 start_gallery.bat watch。",
        "想重新创建桌面快捷方式，双击 CREATE_DESKTOP_SHORTCUT.bat。",
    ):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        set_font(run, size=10.5)

    p = paragraph(doc)
    run = p.add_run("注意：")
    set_font(run, bold=True, color=(160, 70, 30))
    run = p.add_run(" 本项目是本地服务，默认只监听 127.0.0.1:8797。不要把端口暴露到公网或陌生局域网。")
    set_font(run, size=10.5)

    doc.add_heading("2. 包内目录", level=1)
    add_kv_table(
        doc,
        ["路径", "用途", "是否可删除"],
        [
            ("START_GALLERY.bat", "一键启动图库服务。", "不要删"),
            ("CREATE_DESKTOP_SHORTCUT.bat", "为当前用户创建桌面快捷方式。", "可删"),
            ("server.py / web/", "FastAPI 服务和前端页面。", "不要删"),
            ("data/aitag.db", "基础图库数据库，包含作品元数据和咒语。", "不要删"),
            ("data/images/", "基础封面图库。", "不要删"),
            ("data/generated/", "用户自己试生成的图片保存位置。", "可清空"),
            ("data/*.example.json", "空白 token 示例文件。", "可保留"),
            ("logs/", "运行日志。", "可清空"),
        ],
    )

    doc.add_heading("3. Token 与外部 API", level=1)
    paragraph(
        doc,
        "测试包不会包含作者本人的任何 API Key、NovelAI token、Pixiv refresh_token，也不会包含 DeepSeek API Key。所有外部接口都需要测试者自行填写。",
    )
    add_kv_table(
        doc,
        ["功能", "需要的凭据", "保存位置", "填写入口"],
        [
            (
                "NovelAI 试生成",
                "NovelAI Access Token",
                "data/nai_token.local.json",
                "作品详情页右下角工作台设置，或 POST /api/nai/token",
            ),
            (
                "Pixiv 上传/账号",
                "Pixiv refresh_token",
                "data/pixiv_accounts.local.json",
                "/pixiv 页面，账号区域添加 token",
            ),
            (
                "AI 文案/运营分析",
                "用户自己的 OpenAI-compatible API Key",
                "data/ai.local.json",
                "/pixiv 页面，AI Key 区域填写",
            ),
            (
                "DeepSeek",
                "用户自己的 DeepSeek Key",
                "data/ai.local.json",
                "可在 /pixiv 选择 DeepSeek preset 后自行填写",
            ),
        ],
    )
    paragraph(doc, "如果只浏览图库和复制咒语，不需要填写任何 token。没有 token 时，生图和 Pixiv 上传功能会显示未配置或请求失败，这是正常状态。")

    doc.add_heading("4. 本地 HTTP 接口速查", level=1)
    paragraph(doc, "以下接口只服务本机浏览器和测试用户，不建议开放到公网。")
    add_kv_table(
        doc,
        ["接口", "说明"],
        [
            ("GET /", "主图库页面"),
            ("GET /progress", "爬取进度页面"),
            ("GET /generated", "试生成图库页面"),
            ("GET /pixiv", "Pixiv 发布工作台页面"),
            ("GET /api/config", "前端配置"),
            ("GET /api/ai_works_search", "图库搜索，参数 q、prompt、page、page_size、sort、time_range"),
            ("GET /api/work/{id}", "作品详情和多图元数据"),
            ("GET /data/images/{path}", "本地封面图片；缺失时代理 CDN 回退"),
            ("GET /api/plugin/char-swap/extract", "解析 NAI v4 多角色槽位"),
            ("POST /api/plugin/char-swap/transform", "角色替换/克隆"),
            ("POST /api/plugin/char-swap/sanitize", "净化不需要的 tag"),
            ("POST /api/nai/token", "保存用户自己的 NovelAI token"),
            ("POST /api/nai/generate", "串行调用 NovelAI 试生成"),
            ("GET /api/nai/queue", "查看生图队列"),
            ("GET /api/generated", "生成图库分组列表"),
            ("DELETE /api/generated/item/{image_id}", "删除单张生成图"),
            ("GET /api/progress", "爬虫进度 JSON"),
            ("POST /api/crawler/restart", "重启爬虫任务"),
            ("GET /api/pixiv/config", "Pixiv/AI 配置状态，不返回明文 key"),
            ("POST /api/pixiv/accounts", "添加 Pixiv refresh_token"),
            ("POST /api/pixiv/ai-key", "保存用户自己的 AI API Key"),
        ],
    )

    doc.add_heading("5. 常见问题", level=1)
    for question, answer in (
        ("双击后窗口闪退", "先确认安装了 Python 3.11 或更高版本，并且 python 命令在 PATH 中。也可以查看 logs/server.log。"),
        ("网页能打开但图片少量 404", "发布包只内置基础封面图库；详情页多图未缓存时会尝试代理 CDN，离线环境可能打不开。"),
        ("生图失败", "通常是未配置 NovelAI token、token 过期、额度不足或 NovelAI API 网络不可达。"),
        ("Pixiv/AI 文案功能不可用", "发布包不带任何 API key 或 refresh_token，需要用户自己在 /pixiv 页面填写。"),
        ("想完全离线只看图库", "直接使用首页搜索和详情页复制咒语即可，不需要启动爬虫或填写 token。"),
        ("不想留下 token", "删除 data/nai_token.local.json、data/pixiv_accounts.local.json、data/ai.local.json 即可。"),
    ):
        p = paragraph(doc)
        run = p.add_run(question)
        set_font(run, size=10.5, bold=True)
        p = paragraph(doc, answer)
        p.paragraph_format.left_indent = Inches(0.18)

    doc.add_heading("6. 发布包安全说明", level=1)
    for item in (
        "发布包已排除作者本人的 *.local.json 私密文件。",
        "发布包不包含 DeepSeek API Key；DeepSeek 只可能作为用户自行选择的第三方接口名称出现。",
        "发布包默认 AI provider/api_base/model 为空，测试者需要自行配置。",
        "不要把自己的 token 文件再转发给别人。",
    ):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10.5)

    doc.add_heading("7. 本版相对旧 testkit 的修改说明（2026-06-11）", level=1)
    paragraph(
        doc,
        "以下记录从 E:\\aitag-mirror 开发版同步到 E:\\aitag-mirror-gallery 一键包时，相对 2026-06-07 旧 testkit 的主要变更。",
    )
    add_kv_table(
        doc,
        ["类别", "修改内容"],
        [
            ("安装位置", r"一键包目录改为 E:\aitag-mirror-gallery；桌面快捷方式指向该目录的 START_GALLERY.bat。"),
            ("图库数据", "同步最新 data/aitag.db 与 data/images/ 封面图库（约 34658 作品、32936 张本地封面）。"),
            ("新增后端模块", "favorites.py 收藏；ark_char_library.py 明日方舟角色库；nai_prompt_profiles.py 咒语配置；char_marker.py / slot_gender.py 角色槽解析；pixiv_char_tags.py Pixiv 角色 tag。"),
            ("新增数据文件", "data/ark_char_library.json、data/danbooru_style_tags.json。"),
            ("启动脚本", "start_gallery.bat 增加 restart / watch 模式；发布入口仍保留 START_GALLERY.bat 兼容双击启动。"),
            ("配置策略", "gallery_scope=local、gallery_local_only=true、preview_mode=cover_only：本地只存封面，详情多图走 CDN 回退。"),
            ("前端", "web/app.js 等更新：收藏页 /favorites、角色库搜索、char-swap 插件等。"),
            ("依赖", "requirements.txt 增加打码后处理相关依赖（gradio、ultralytics、opencv 等）。"),
            ("安全剔除", "不打包 nai_token.local.json、pixiv_accounts.local.json、ai.local.json、pixiv_chrome_profiles/、logs/ 历史日志。"),
            ("开发目录保留", r"E:\aitag-mirror 仍为开发/爬虫工作区；本包为只读分发副本，日常浏览请用本目录。"),
        ],
    )
    paragraph(doc, "如需继续爬取或改代码，请在 E:\\aitag-mirror 开发目录操作，完成后再重新运行 scripts\\make_release.ps1 刷新本包。")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("aitag-mirror local testkit")
    set_font(run, size=8, color=(120, 120, 120))

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
